import flet as ft
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import re
from morphology import replace_variables_with_case, change_case, clean_text

# Функция для замены переменных в тексте
def replace_variables_in_text(text, entries, optional_entries):
    """
    Заменяет обязательные и опциональные переменные в тексте.
    """

    # Замена обязательных полей
    variables = {
        "plaintiff": entries['plaintiff'].strip() if isinstance(entries['plaintiff'], str) else entries[
            'plaintiff'].value.strip(),
        "defendant": entries['defendant'].strip() if isinstance(entries['defendant'], str) else entries[
            'defendant'].value.strip(),
        "decision_number": entries['decision_number'].strip() if isinstance(entries['decision_number'], str) else
        entries['decision_number'].value.strip(),
        "uid": entries['uid'].strip() if isinstance(entries['uid'], str) else entries['uid'].value.strip(),
        "decision_date": entries['decision_date'].strip() if isinstance(entries['decision_date'], str) else entries[
            'decision_date'].value.strip(),
        "crux_of_the_matter": entries['crux'].strip() if isinstance(entries['crux'], str) else entries[
            'crux'].value.strip(),
        "third_faces": entries['third_faces'].strip() if isinstance(entries['third_faces'], str) else entries[
            'third_faces'].value.strip(),
        "applicant": entries['applicant'].strip() if isinstance(entries['applicant'], str) else entries[
            'applicant'].value.strip(),
        "interested_persons": entries['interested_persons'].strip() if isinstance(entries['interested_persons'], str) else entries[
            'interested_persons'].value.strip(),
    }

    # Сначала заменяем обязательные переменные
    for key, value in variables.items():
        text = re.sub(r"\{" + re.escape(key) + r"\}", value, text)

    # Теперь заменяем переменные с падежами
    text = replace_variables_with_case(text, variables)

    # Замена опциональных переменных
    for key, entry in optional_entries.items():
        # Если entry — строка, используем её напрямую
        if isinstance(entry, str):
            optional_value = entry.strip()
        # Если entry — контейнер, пытаемся извлечь значение из TextField
        elif isinstance(entry, ft.Container):
            # Проверяем, является ли контент контейнера TextField
            if isinstance(entry.content, ft.TextField):
                optional_value = entry.content.value.strip() if entry.content.value else ""
            else:
                optional_value = ""  # Если в контейнере нет TextField, то оставляем пустое значение
        else:
            optional_value = ""  # В случае других типов (например, если это не контейнер)

        # Заменяем опциональные поля в тексте
        if optional_value:
            text = re.sub(r"\{" + re.escape(key) + r"\}", optional_value, text)
        else:
            text = re.sub(r"\{" + re.escape(key) + r"\}", "", text)  # Если нет значения, удаляем

    return text

#Функция для создания обычного решения суда
def create_document(entries, descriptive_text, pleading_text, optional_entries, plaintiff_attendance_var, defendant_attendance_var, third_party_attendance_var, selected_template, page):
    try:
        # Получение данных из полей
        decision_number = entries['decision_number'].value.strip()
        uid = entries['uid'].value.strip()
        decision_date = entries['decision_date'].value.strip()
        plaintiff = entries['plaintiff'].value.strip()
        defendant = entries['defendant'].value.strip()
        third_faces = entries['third_faces'].value.strip()
        crux_of_the_matter = entries['crux'].value.strip()
        descriptive_part = descriptive_text.value.strip()
        pleading_part = pleading_text.value.strip()

        # Проверка явки сторон
        plaintiff_attendance = plaintiff_attendance_var if isinstance(plaintiff_attendance_var, str) else plaintiff_attendance_var.value
        defendant_attendance = defendant_attendance_var if isinstance(defendant_attendance_var, str) else defendant_attendance_var.value
        third_party_attendance = third_party_attendance_var if isinstance(third_party_attendance_var, str) else third_party_attendance_var.value

        # Создание документа
        doc = docx.Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        # Установка полей документа
        section = doc.sections[0]
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        # Настройка междустрочного интервала
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_after = 0  # Убираем интервал после абзаца

        # Заполнение документа
        doc.add_paragraph(decision_number).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph(uid).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph("РЕШЕНИЕ \n Именем Российской Федерации", style='Normal').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = decision_date
        table.cell(0, 1).text = "г. Оренбург"
        table.cell(0, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Определение участников
        participants = []
        if plaintiff_attendance == "Да":
            participants.append(f"истца {change_case(plaintiff, 'gent')}")
        if defendant_attendance == "Да":
            participants.append(f"ответчика {change_case(defendant, 'gent')}")
        if third_party_attendance == "Да":
            participants.append("третьих лиц")

        participants_text = f"\nс участием {', '.join(participants)}" if participants else ""

        doc.add_paragraph(
            f'Оренбургский районный суд Оренбургской области в составе\nпредседательствующего судьи Мичуриной Т.А.,\nпри секретаре Васильевой Е.Ю.,{participants_text},'.strip()
        )

        plaintiff_genitive = change_case(plaintiff, 'gent')
        defendant_dative = change_case(defendant, 'datv')

        doc.add_paragraph(
            f'рассмотрев в открытом судебном заседании гражданское дело по исковому заявлению {plaintiff_genitive} к {defendant_dative} {crux_of_the_matter},',
            style='Normal'
        ).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        doc.add_paragraph("УСТАНОВИЛ:").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        #Чистим текст от лишнего
        descriptive_part = clean_text(descriptive_part)
        pleading_part = clean_text(pleading_part)

        #Добавляем описательную часть
        if descriptive_part:
            paragraph = doc.add_paragraph(f"Истец обратился в суд с вышеуказанным иском, указав, что {descriptive_part}", style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        #Добавлять просительную часть
        if pleading_part:
            paragraph = doc.add_paragraph(f'Просит суд {pleading_part}', style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        # Проверка явки сторон
        if plaintiff_attendance == "Да":
            paragraph = doc.add_paragraph(f"Истец {plaintiff} в судебном заседании исковые требования поддержал",
                                          style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0
        else:
            paragraph = doc.add_paragraph(
                f"Истец {plaintiff} в судебное заседание не явился, извещен надлежащим образом.", style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        if defendant_attendance == "Да":
            paragraph = doc.add_paragraph(
                f"Ответчик {defendant} в судебном заседании возражал против удовлетворения исковых требований, пояснил, что",
                style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0
        else:
            paragraph = doc.add_paragraph(
                f"Ответчик {defendant} в судебное заседание не явился, извещен надлежащим образом.", style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        # Проверка явки третьего лица
        if third_party_attendance == "Да":
            paragraph = doc.add_paragraph(f"Третье лицо {third_faces} в судебном заседании поддержало требования истца",
                                          style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0
        else:
            paragraph = doc.add_paragraph(
                f"Третье лицо {third_faces} в судебное заседание не явилось, извещено надлежащим образом.",
                style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)

        paragraph = doc.add_paragraph(
            "Суд определил рассмотреть дело в отсутствие не явившихся лиц, в порядке ст. 167  ГПК Российской Федерации.",
            style='Normal')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(1.25)

        # Вставка текста из выбранного шаблона
        if selected_template:
            template_path = os.path.join("Templates of motivational parts", selected_template)
            if os.path.exists(template_path):
                template_doc = docx.Document(template_path)

                for paragraph in template_doc.paragraphs:
                    # Добавляем новый абзац в документ
                    new_paragraph = doc.add_paragraph(style='Normal')
                    new_paragraph.alignment = paragraph.alignment
                    new_paragraph.paragraph_format.first_line_indent = Cm(1.25)
                    new_paragraph.paragraph_format.space_after = 0

                    # Получаем текст из абзаца и заменяем переменные
                    text = paragraph.text

                    # Заменяем обязательные и опциональные переменные
                    text = replace_variables_in_text(text, entries, optional_entries)

                    # Перезаписываем текст в абзаце
                    new_paragraph.add_run(text)

        # Укажите путь к папке, куда сохранять документы
        save_folder = r"C:\Python Projects\WordCreator\Court Document Creator\Saved Docs"

        # Проверим, существует ли папка, если нет — создадим её
        if not os.path.exists(save_folder):
            try:
                os.makedirs(save_folder)
                print(f"Папка {save_folder} была создана.")
            except Exception as e:
                print(f"Не удалось создать папку: {e}")
                return  # Завершаем выполнение функции, если не удалось создать папку

        # Пример имени файла, заменим символы на допустимые
        filename = f"{decision_number} {plaintiff} {crux_of_the_matter}.docx"

        # Заменим символы, которые могут вызвать ошибку
        filename = filename.replace("/", "_").replace("\\", "_").replace(":", "_").replace("*", "_").replace("?",
                                                                                                             "_").replace(
            "\"", "_").replace("<", "_").replace(">", "_").replace("|", "_")

        # Полный путь к файлу
        file_path = os.path.join(save_folder, filename)

        # Сохраняем документ
        doc.save(file_path)

        # Уведомляем пользователя о сохранении
        page.snack_bar = ft.SnackBar(ft.Text(f"Документ успешно сохранен: {file_path}"))
        page.snack_bar.open = True
        page.update()

    except Exception as e:
        print(f"Произошла ошибка: {e}")
        page.snack_bar = ft.SnackBar(ft.Text("Произошла ошибка при создании документа."))
        page.snack_bar.open = True
        page.update()

#Функция для создания заочного решения суда
def create_extramural_document(entries, descriptive_text, pleading_text, optional_entries, plaintiff_attendance_var, defendant_attendance_var, third_party_attendance_var, selected_template, page):
    try:

        # Получение данных из полей
        decision_number = entries['decision_number'].value.strip()
        uid = entries['uid'].value.strip()
        decision_date = entries['decision_date'].value.strip()
        plaintiff = entries['plaintiff'].value.strip()
        defendant = entries['defendant'].value.strip()
        third_faces = entries['third_faces'].value.strip()
        crux_of_the_matter = entries['crux'].value.strip()
        descriptive_part = descriptive_text.value.strip()
        pleading_part = pleading_text.value.strip()


        # Проверка явки сторон
        plaintiff_attendance = plaintiff_attendance_var if isinstance(plaintiff_attendance_var,str) else plaintiff_attendance_var.value
        defendant_attendance = defendant_attendance_var if isinstance(defendant_attendance_var,str) else defendant_attendance_var.value
        third_party_attendance = third_party_attendance_var if isinstance(third_party_attendance_var,str) else third_party_attendance_var.value

        # Создание документа
        doc = docx.Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        # Установка полей документа
        section = doc.sections[0]
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        # Настройка междустрочного интервала
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_after = 0  # Убираем интервал после абзаца

        # Заполнение документа
        doc.add_paragraph(decision_number).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph(uid).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph("ЗАОЧНОЕ РЕШЕНИЕ \n Именем Российской Федерации",
                          style='Normal').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = decision_date
        table.cell(0, 1).text = "г. Оренбург"
        table.cell(0, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Определение участников
        participants = []
        if plaintiff_attendance == "Да":
            participants.append(f"истца {change_case(plaintiff, 'gent')}")
        if defendant_attendance == "Да":
            participants.append(f"ответчика {change_case(defendant, 'gent')}")
        if third_party_attendance == "Да":
            participants.append("третьих лиц")

        participants_text = f"\nс участием {', '.join(participants)}" if participants else ""

        doc.add_paragraph(
            f'Оренбургский районный суд Оренбургской области в составе\nпредседательствующего судьи Мичуриной Т.А.,\nпри секретаре Васильевой Е.Ю.,{participants_text},'.strip()
        )

        plaintiff_genitive = change_case(plaintiff, 'gent')
        defendant_dative = change_case(defendant, 'datv')

        doc.add_paragraph(
            f'рассмотрев в открытом судебном заседании гражданское дело по исковому заявлению {plaintiff_genitive} к {defendant_dative} {crux_of_the_matter},',
            style='Normal'
        ).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        doc.add_paragraph("УСТАНОВИЛ:").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Чистим текст от лишнего
        descriptive_part = clean_text(descriptive_part)
        pleading_part = clean_text(pleading_part)

        # Добавляем описательную часть
        if descriptive_part:
            paragraph = doc.add_paragraph(
                f"Истец обратился в суд с вышеуказанным иском, указав, что {descriptive_part}", style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        # Добавлять просительную часть
        if pleading_part:
            paragraph = doc.add_paragraph(f'Просит суд {pleading_part}', style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        # Проверка явки сторон
        if plaintiff_attendance == "Да":
            paragraph = doc.add_paragraph(f"Истец {plaintiff} в судебном заседании исковые требования поддержал",
                                          style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0
        else:
            paragraph = doc.add_paragraph(
                f"Истец {plaintiff} в судебное заседание не явился, извещен надлежащим образом.", style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        # Проверка явки третьего лица
        if third_party_attendance == "Да":
            paragraph = doc.add_paragraph(f"Третье лицо {third_faces} в судебном заседании поддержало требования истца",
                                          style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0
        else:
            paragraph = doc.add_paragraph(
                f"Третье лицо {third_faces} в судебное заседание не явилось, извещено надлежащим образом.",
                style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)

        # Проверка явки ответчика
        if defendant_attendance == "Да":
            paragraph = doc.add_paragraph(
                f"Ответчик {defendant} в судебном заседании возражал против удовлетворения исковых требований, пояснил, что",
                style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0
        else:
            paragraph = doc.add_paragraph(
                f"Ответчик {defendant} в судебное заседание не явился, извещался судом заблаговременно о времени и месте судебного разбирательства надлежащим образом."
                f"Согласно сведениям адресно-справочного бюро ОУФМС России по Оренбургской области, {defendant} зарегистрирован по адресу: ."
                "Корреспонденция направлялась ответчику по адресу регистрации и фактическому месту проживания, конверты возвращены в суд с отметкой: «истек срок хранения»."
                "Руководствуясь статьями 167, 233 Гражданского процессуального кодекса Российской Федерации, суд определил рассмотреть дело в отсутствие неявившихся сторон в порядке заочного производства",
                style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        # Вставка текста из выбранного шаблона
        if selected_template:
            template_path = os.path.join("Templates of motivational parts", selected_template)
            if os.path.exists(template_path):
                template_doc = docx.Document(template_path)

                for paragraph in template_doc.paragraphs:
                    # Добавляем новый абзац в документ
                    new_paragraph = doc.add_paragraph(style='Normal')
                    new_paragraph.alignment = paragraph.alignment
                    new_paragraph.paragraph_format.first_line_indent = Cm(1.25)
                    new_paragraph.paragraph_format.space_after = 0

                    # Получаем текст из абзаца и заменяем переменные
                    text = paragraph.text

                    # Заменяем обязательные и опциональные переменные
                    text = replace_variables_in_text(text, entries, optional_entries)

                    # Перезаписываем текст в абзаце
                    new_paragraph.add_run(text)

        #Обжалование
        paragraph = doc.add_paragraph(
            "Ответчик вправе подать в суд, принявший заочное решение, заявление об отмене этого решения суда в течение семи дней со дня вручения ему копии этого решения. "
            "Ответчиком заочное решение суда может быть обжаловано в апелляционном порядке в Оренбургский областной суд через Оренбургский районный суд Оренбургской области в течение одного месяца со дня вынесения определения суда об отказе в удовлетворении заявления об отмене этого решения суда. "
            "Иными лицами, участвующими в деле, а также лицами, которые не были привлечены к участию в деле и вопрос о правах и об обязанностях которых был разрешен судом, заочное решение суда может быть обжаловано в апелляционном порядке в Оренбургский областной суд через Оренбургский районный суд Оренбургской области "
            "в течение одного месяца по истечении срока подачи ответчиком заявления об отмене этого решения суда, а в случае, если такое заявление подано, - в течение одного месяца со дня вынесения определения суда об отказе в удовлетворении этого заявления. "
            "Заочное решение принято в окончательной форме  ",
            style='Normal')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(1.25)

        # Укажите путь к папке, куда сохранять документы
        save_folder = r"C:\Python Projects\WordCreator\Court Document Creator\Saved Docs"

        # Проверим, существует ли папка, если нет — создадим её
        if not os.path.exists(save_folder):
            try:
                os.makedirs(save_folder)
                print(f"Папка {save_folder} была создана.")
            except Exception as e:
                print(f"Не удалось создать папку: {e}")
                return  # Завершаем выполнение функции, если не удалось создать папку

        # Пример имени файла, заменим символы на допустимые
        filename = f"{decision_number} {plaintiff} {crux_of_the_matter}.docx"

        # Заменим символы, которые могут вызвать ошибку
        filename = filename.replace("/", "_").replace("\\", "_").replace(":", "_").replace("*", "_").replace("?",
                                                                                                             "_").replace(
            "\"", "_").replace("<", "_").replace(">", "_").replace("|", "_")

        # Полный путь к файлу
        file_path = os.path.join(save_folder, filename)

        # Сохраняем документ
        doc.save(file_path)

        # Уведомляем пользователя о сохранении
        page.snack_bar = ft.SnackBar(ft.Text(f"Документ успешно сохранен: {file_path}"))
        page.snack_bar.open = True
        page.update()

    except Exception as e:
        print(f"Произошла ошибка: {e}")
        page.snack_bar = ft.SnackBar(ft.Text("Произошла ошибка при создании документа."))
        page.snack_bar.open = True
        page.update()

#Функция для создания решения суда в особом порядке
def create_special_document(entries, descriptive_text, pleading_text, optional_entries, applicant_attendance_var, interested_persons_attendance_var, selected_template, page):
    try:

        # Получение данных из полей
        decision_number = entries['decision_number'].value.strip()
        uid = entries['uid'].value.strip()
        decision_date = entries['decision_date'].value.strip()
        applicant = entries['applicant'].value.strip()
        interested_persons = entries['interested_persons'].value.strip()
        crux_of_the_matter = entries['crux'].value.strip()
        descriptive_part = descriptive_text.value.strip()
        pleading_part = pleading_text.value.strip()

        # Проверка явки сторон
        applicant_attendance = applicant_attendance_var if isinstance(applicant_attendance_var, str) else applicant_attendance_var.value
        interested_persons_attendance = interested_persons_attendance_var if isinstance(interested_persons_attendance_var, str) else interested_persons_attendance_var.value

        # Создание документа
        doc = docx.Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        # Установка полей документа
        section = doc.sections[0]
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        # Настройка междустрочного интервала
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_after = 0  # Убираем интервал после абзаца

        # Заполнение документа
        doc.add_paragraph(decision_number).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph(uid).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph("РЕШЕНИЕ \n Именем Российской Федерации", style='Normal').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = decision_date
        table.cell(0, 1).text = "г. Оренбург"
        table.cell(0, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Определение участников
        participants = []
        if applicant_attendance == "Да":
            participants.append(f"заявителя {change_case(applicant, 'gent')}")
        if interested_persons == "Да":
            participants.append(f"заинтересованного лица {change_case(interested_persons, 'gent')}")

        participants_text = f"\nс участием {', '.join(participants)}" if participants else ""

        doc.add_paragraph(
            f'Оренбургский районный суд Оренбургской области в составе\nпредседательствующего судьи Мичуриной Т.А.,\nпри секретаре Васильевой Е.Ю.,{participants_text},'.strip()
        )

        applicant_genitive = change_case(applicant, 'gent')

        doc.add_paragraph(
            f'рассмотрев в открытом судебном заседании гражданское дело по заявлению {applicant_genitive} {crux_of_the_matter},',
            style='Normal'
        ).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        doc.add_paragraph("УСТАНОВИЛ:").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        #Чистим текст от лишнего
        descriptive_part = clean_text(descriptive_part)
        pleading_part = clean_text(pleading_part)

        #Добавляем описательную часть
        if descriptive_part:
            paragraph = doc.add_paragraph(f"Заявитель обратился в суд с вышеуказанным иском, указав, что {descriptive_part}", style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        #Добавлять просительную часть
        if pleading_part:
            paragraph = doc.add_paragraph(f'Просит суд {pleading_part}', style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        # Проверка явки сторон
        if applicant_attendance == "Да":
            paragraph = doc.add_paragraph(f"Заявитель {applicant} в судебном заседании заявленные требования поддержал",
                                          style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0
        else:
            paragraph = doc.add_paragraph(
                f"Заявитель {applicant} в судебное заседание не явился, извещен надлежащим образом.", style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        if interested_persons_attendance == "Да":
            paragraph = doc.add_paragraph(
                f"Заинтересованное лицо {interested_persons} в судебном заседании пояснил, что",
                style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0
        else:
            paragraph = doc.add_paragraph(
                f"Заинтересованное лицо {interested_persons} в судебное заседание не явился, извещен надлежащим образом.", style='Normal')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = 0

        paragraph = doc.add_paragraph(
            "Суд определил рассмотреть дело в отсутствие не явившихся лиц, в порядке ст. 167  ГПК Российской Федерации.",
            style='Normal')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(1.25)

        # Вставка текста из выбранного шаблона
        if selected_template:
            template_path = os.path.join("Templates of motivational parts", selected_template)
            if os.path.exists(template_path):
                template_doc = docx.Document(template_path)

                for paragraph in template_doc.paragraphs:
                    # Добавляем новый абзац в документ
                    new_paragraph = doc.add_paragraph(style='Normal')
                    new_paragraph.alignment = paragraph.alignment
                    new_paragraph.paragraph_format.first_line_indent = Cm(1.25)
                    new_paragraph.paragraph_format.space_after = 0

                    # Получаем текст из абзаца и заменяем переменные
                    text = paragraph.text

                    # Заменяем обязательные и опциональные переменные
                    text = replace_variables_in_text(text, entries, optional_entries)

                    # Перезаписываем текст в абзаце
                    new_paragraph.add_run(text)

        # Укажите путь к папке, куда сохранять документы
        save_folder = r"C:\Python Projects\WordCreator\Court Document Creator\Saved Docs"

        # Проверим, существует ли папка, если нет — создадим её
        if not os.path.exists(save_folder):
            try:
                os.makedirs(save_folder)
                print(f"Папка {save_folder} была создана.")
            except Exception as e:
                print(f"Не удалось создать папку: {e}")
                return  # Завершаем выполнение функции, если не удалось создать папку

        # Пример имени файла, заменим символы на допустимые
        filename = f"{decision_number} {applicant} {crux_of_the_matter}.docx"

        # Заменим символы, которые могут вызвать ошибку
        filename = filename.replace("/", "_").replace("\\", "_").replace(":", "_").replace("*", "_").replace("?",
                                                                                                             "_").replace(
            "\"", "_").replace("<", "_").replace(">", "_").replace("|", "_")

        # Полный путь к файлу
        file_path = os.path.join(save_folder, filename)

        # Сохраняем документ
        doc.save(file_path)

        # Уведомляем пользователя о сохранении
        page.snack_bar = ft.SnackBar(ft.Text(f"Документ успешно сохранен: {file_path}"))
        page.snack_bar.open = True
        page.update()

    except Exception as e:
        print(f"Произошла ошибка: {e}")
        page.snack_bar = ft.SnackBar(ft.Text("Произошла ошибка при создании документа."))
        page.snack_bar.open = True
        page.update()